#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Fri Nov  9 10:07:11 2018

@author: joaquim
"""
from __future__ import print_function
import requests
import json#, datetime
#import os
import ssl
import time
from docx import Document
#from docx.shared import Inches

total_time = 0
start = time.time() 

print(ssl.OPENSSL_VERSION)
try:
  from ConfigParser import ConfigParser
except ImportError:
  from configparser import ConfigParser

  
global field_ticket, username, password, target_projects
# get info from config.ini
config = ConfigParser()
try:
  config.read('config.ini')
except Exception:
  print("error reading config.ini")


#Get fields from config.ini file
username = config.get("Field", "username")
password = config.get("Field", "password")
client_id = config.get("HQ", "client_id")
client_secret = config.get("HQ", "client_secret")
account_id = config.get("HQ", "account_id")

#Set target projects
target_projects = ['SITEUR - MANTENIMIENTO TUNEL']


global base_url_field
base_url_field = "https://bim360field.eu.autodesk.com/"

all_project_hqs = []
all_project_ids = []
all_project_names = []


# control what info you want to download
harvest_info = {}
print("Downloading:")
for name in config.options('download'):
  harvest_info[name] = config.getboolean('download',name)
  print("{:20}: {}".format(name,config.getboolean('download',name)))


# function for FIELD commands
field_ticket = ""
def field_api_cmd(method, command, payload):
  """
  This function is for calling all commands from the Field cloud
  It will check first if there is a ticket available otherwise it will login and get a ticket first.
  When calling the function you can give it the various payload options that are documented by Autodesk
  """
  global field_ticket
  if field_ticket == "":
    #if not logedin yet, login
    field_ticket = json.loads(
      requests.request(
        "POST","{0}/api/login".format(base_url_field),
        data = {
          "username": username,
          "password": password
        },
        headers = {
          'content-type': "application/x-www-form-urlencoded",
          'Cache-Control': "no-cache"
        }).text)['ticket']
    print("Logging into {0}".format(base_url_field))
    print('Field ticket: {0}'.format(field_ticket))

  url = '{base_url}{command}'.format(base_url=base_url_field, command=command)
  global headers_field
  headers = {
    'content-type': "application/x-www-form-urlencoded",
    'Cache-Control': "no-cache",
  }
  headers_field = headers
  
  global payload_field
  payload['ticket'] = field_ticket
  payload_field = payload
  #print("FIELD Request= method:{method}, url:{url}".format(method=method,url=url))
  #print("data:{}".format(payload))
  #print("headers:{}".format(headers))
  
  response = requests.request(method, url, data=payload, headers=headers).json()  
  print("FIELD returned: {} records".format(len(response)))
  return response


# get all field projects from field
get_projects = field_api_cmd("POST", "api/projects",{})

#get project ID's
all_project_ids = []
for i in get_projects:
    if i['name'] in target_projects:
        all_project_ids.append(i['project_id'])
        print('Adding project, name: {}, ID: {}'.format(i['name'], i['project_id']))


"""
    
# standard function for getting info from field
def get_standard_field_records(type_name, method, command, **payload):
  all_records = []  
  for i, current_project_id in enumerate(all_project_ids):
    if 'project_id' in payload:
      payload['project_id'] = current_project_id
    records = field_api_cmd(method, command, payload)
    if records:
      for record in records:
        record['project_id'] = current_project_id
      print('Adding {} {} for project: {}'.format(len(records), type_name, all_project_names[i]))
      all_records.append(records)
    else:
      print('{} {} for project: {}'.format(len(records), type_name, all_project_names[i]))
    double_save_file("{}_info".format(type_name), all_records)

"""


def get_checklists(project_id):
    parameters = {}
    parameters['project_id'] = project_id
    #parameters['project_id'] = all_project_ids[0]
    parameters['limit'] = 100
    global project_checklists
    project_checklists = field_api_cmd("GET", "fieldapi/checklists/v1", parameters)
    global checklist_0
    checklist_0 = project_checklists[1]['id']
    
def get_checklist_details(id, project_id):
    parameters = {}
    parameters['project_id'] = project_id
    parameters['id'] = id
    global checklist
    checklist = field_api_cmd("GET", "fieldapi/checklists/v1/{}".format(id), parameters)
    build_report(checklist)
    
    
def build_report(checklist):
    #create document
    document = Document()       
    #add heading to document
    document.add_heading(checklist['name'],0)  
    #add intro paragraph
    document.add_paragraph('Informe asociado al checklist.')
    
    for section in checklist['sections']:
        if(len(section['section_name']) > 0):           
            #create Heading with section name
            document.add_heading(section['section_name'], level=1)                      
            for item in section['items']:
                if(len(item['response']) == 0):
                    item['response'] = 'Sin respuesta'
                #create numbered list containing every item in section    
                document.add_paragraph(
                    item['question_text']+': '+item['response'], style='List Number'        
                )                              
        
    save_string = checklist['name']+checklist['identifier'] + '.docx'      
    #add page break
    document.add_page_break()
    #save document
    document.save(save_string)
       
            
#Run get_checklists   
for project in all_project_ids:
    get_checklists(project)
    for c_list in project_checklists:
        get_checklist_details(c_list['id'], project)


end = time.time()
total_time = round(end - start, 2)
print('All reports generated successfully. Total execution time: {} seconds'.format(total_time))

